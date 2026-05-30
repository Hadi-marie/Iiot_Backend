from __future__ import annotations

import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path

import docx

from .config import settings

# A category header in the source doc looks like "1️⃣ أسئلة عامة عن الخدمة".
# Those keycap emoji are: ASCII digit + U+FE0F (variation selector) + U+20E3 (keycap).
_KEYCAP = "\u20e3"
# A real question line looks like "12. هل يوجد دعم تقني؟"
_QUESTION_RE = re.compile(r"^\s*(\d{1,3})\s*[\.\)]\s+(.+)$")
_ANSWER_MARKER_RE = re.compile(r"^\s*(الجواب|الإجابة|answer)\s*[:：]?\s*", re.IGNORECASE)


@dataclass
class QAPair:
    qid: int
    question: str
    answer: str
    category: str


@dataclass
class KnowledgeDocument:
    doc_id: str
    text: str          # text used to build the embedding vector
    answer: str        # text returned to the user
    doc_type: str      # "qa" | "passage" | "table"
    source: str
    question: str = ""
    category: str = ""
    metadata: dict = field(default_factory=dict)


def _clean(text: str) -> str:
    text = text.replace("\u200f", "").replace("\u200e", "")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _is_category_header(line: str) -> bool:
    return _KEYCAP in line


def _is_answer_marker(line: str) -> bool:
    return bool(re.match(r"^\s*(الجواب|الإجابة|answer)\b", line, re.IGNORECASE))


def parse_answers_docx(path: str | None = None) -> list[QAPair]:
    """Parse Answers_bot.docx into an ordered list of curated question/answer pairs.

    A numbered line is treated as a *real* question only when the next non-empty
    paragraph is an answer marker ("الجواب"). This prevents numbered sub-lists that
    appear inside an answer (e.g. "1. Network IDS", "2. Host IDS") from being mistaken
    for new questions.
    """
    docx_path = Path(path or settings.knowledge_docx)
    if not docx_path.exists():
        raise FileNotFoundError(f"Knowledge document not found: {docx_path}")

    document = docx.Document(str(docx_path))
    lines = [_clean(p.text) for p in document.paragraphs]
    lines = [ln for ln in lines if ln]

    def next_nonempty(idx: int) -> str | None:
        return lines[idx + 1] if idx + 1 < len(lines) else None

    pairs: list[QAPair] = []
    current_category = "عام"
    current: QAPair | None = None
    answer_lines: list[str] = []

    def flush() -> None:
        nonlocal current, answer_lines
        if current is not None:
            answer = _clean("\n".join(answer_lines))
            if answer:
                current.answer = answer
                pairs.append(current)
        current = None
        answer_lines = []

    for idx, line in enumerate(lines):
        if _is_category_header(line):
            flush()
            current_category = re.sub(r"^\W*\d+\W*", "", line).strip() or current_category
            continue

        match = _QUESTION_RE.match(line)
        following = next_nonempty(idx)
        is_real_question = bool(match) and following is not None and _is_answer_marker(following)

        if is_real_question:
            flush()
            qid = int(match.group(1))
            question = _clean(match.group(2))
            current = QAPair(qid=qid, question=question, answer="", category=current_category)
            continue

        if current is not None:
            if _is_answer_marker(line):
                stripped = _ANSWER_MARKER_RE.sub("", line)
                if stripped.strip():
                    answer_lines.append(stripped)
                continue
            answer_lines.append(line)

    flush()

    # Deduplicate by question id, keeping the richest answer.
    best: dict[int, QAPair] = {}
    for pair in pairs:
        existing = best.get(pair.qid)
        if existing is None or len(pair.answer) > len(existing.answer):
            best[pair.qid] = pair
    return [best[k] for k in sorted(best)]


def extract_docx_tables(path: str | None = None) -> list[str]:
    """Comparison tables (IDS vs IPS, pricing tiers, etc.) as readable text blocks."""
    docx_path = Path(path or settings.knowledge_docx)
    document = docx.Document(str(docx_path))
    blocks: list[str] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [_clean(c.text) for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if len(rows) >= 2:
            blocks.append("\n".join(rows))
    return blocks


def _chunk_text(text: str, max_chars: int = 900, overlap: int = 120) -> list[str]:
    text = _clean(text)
    if len(text) <= max_chars:
        return [text] if text else []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        # Try to break on a sentence boundary near the end.
        window = text[start:end]
        cut = max(window.rfind("\n"), window.rfind(". "), window.rfind("؟ "), window.rfind("، "))
        if cut > max_chars * 0.5 and end < len(text):
            end = start + cut + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = max(end - overlap, end) if end >= len(text) else end - overlap
    return chunks


def load_markdown_docs(directory: str | None = None) -> list[tuple[str, str]]:
    """Return (source_name, text) for markdown knowledge files."""
    target = directory or settings.extra_knowledge_dir
    out: list[tuple[str, str]] = []
    if not target:
        return out
    base = Path(target)
    if not base.exists():
        return out
    for md in sorted(base.rglob("*.md")):
        try:
            out.append((md.name, md.read_text(encoding="utf-8", errors="ignore")))
        except OSError:
            continue
    return out


_DOC_NS = uuid.UUID("6f0c8a9e-1c2b-4d3e-8f4a-2b1c3d4e5f60")


def _doc_id(*parts: str) -> str:
    return str(uuid.uuid5(_DOC_NS, "::".join(parts)))


def build_documents() -> list[KnowledgeDocument]:
    """Build every embeddable document: curated Q&A, answer passages, tables, and project docs."""
    docs: list[KnowledgeDocument] = []
    docx_name = Path(settings.knowledge_docx).name

    pairs = parse_answers_docx()
    for pair in pairs:
        # 1) Authoritative Q&A: the embedding is built from the question so a user
        #    question matches it directly. The exact curated answer is returned.
        docs.append(
            KnowledgeDocument(
                doc_id=_doc_id("qa", str(pair.qid), pair.question),
                text=f"{pair.question}\n{pair.category}",
                answer=pair.answer,
                doc_type="qa",
                source=docx_name,
                question=pair.question,
                category=pair.category,
                metadata={"qid": pair.qid},
            )
        )
        # 2) Answer passages: let semantic search also hit the answer body for
        #    topic questions that are phrased very differently.
        for idx, chunk in enumerate(_chunk_text(pair.answer)):
            docs.append(
                KnowledgeDocument(
                    doc_id=_doc_id("passage", str(pair.qid), str(idx)),
                    text=chunk,
                    answer=chunk,
                    doc_type="passage",
                    source=docx_name,
                    question=pair.question,
                    category=pair.category,
                    metadata={"qid": pair.qid, "chunk": idx},
                )
            )

    for t_idx, block in enumerate(extract_docx_tables()):
        docs.append(
            KnowledgeDocument(
                doc_id=_doc_id("table", str(t_idx)),
                text=block,
                answer=block,
                doc_type="table",
                source=docx_name,
                category="جداول مقارنة",
                metadata={"table": t_idx},
            )
        )

    for source_name, text in load_markdown_docs():
        for idx, chunk in enumerate(_chunk_text(text, max_chars=1100, overlap=150)):
            docs.append(
                KnowledgeDocument(
                    doc_id=_doc_id("md", source_name, str(idx)),
                    text=chunk,
                    answer=chunk,
                    doc_type="passage",
                    source=source_name,
                    metadata={"chunk": idx},
                )
            )

    return docs
